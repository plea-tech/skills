#!/usr/bin/env python3
"""
从 report_draft.md 生成符合 output_define 样式的竞品分析报告 docx。
Step 3 迭代1 的「整理为 docx」步骤：以草稿为唯一正文来源，应用文档风格定义。
字体、字号、段落格式严格参照 .cursor/skills/Competitor_Analysis_Report_d1/references/output_define.md。
约定：在项目根目录下执行，reports/、extracted_content/ 均相对于当前工作目录。
"""
import re
import sys
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_BREAK

# 与 output_define.md §2 字体与字号、§3 标题与段落 一致
FONT_HEADING_ZH = "微软雅黑"
FONT_HEADING_EN = "Arial"
FONT_BODY_ZH = "微软雅黑"
FONT_BODY_EN = "Calibri"
FONT_REF_ZH = "宋体"
FONT_REF_EN = "Times New Roman"
PT_HEADING1 = 16   # 16–18pt
PT_HEADING2 = 14   # 14pt
PT_HEADING3 = 12   # 12pt 三级标题
PT_BODY = 10.5     # 五号
PT_SUBHEADING = 11    # 小标题（如「本地文件：」）
PT_COVER_TITLE = 24    # 封面主标题 22–26pt
PT_COVER_SUB = 12      # 封面副标题 小四
SPACE_BEFORE_H1 = 24   # 24–30pt
SPACE_AFTER_H1 = 12    # 12pt
SPACE_BEFORE_H2 = 14   # 12–18pt
SPACE_AFTER_H2 = 6    # 6pt
SPACE_BEFORE_H3 = 10   # 6–12pt
SPACE_AFTER_H3 = 3    # 3pt
FIRST_LINE_INDENT_CM = 0.74  # 首行缩进 2 字符
LINE_SPACING_BODY = 1.25    # 1.25–1.5 倍


def _project_root() -> Path:
    """项目目录（约定：脚本在项目根下执行，即 cwd）。"""
    return Path.cwd()


def _apply_page_setup(doc):
    """output_define §1 页面与版心：A4，上下 2.54cm、左右 3.17cm。"""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def _add_cover_page(doc, title: str, analyst: str, report_date: str):
    """output_define §7 封面（必选，单独一页）：主标题居中 22–26pt，副标题含分析人、分析时间，封面后分页。"""
    # 主标题
    p0 = doc.add_paragraph()
    p0.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(18)
    p0.paragraph_format.line_spacing = 1.5  # 1.5 倍行距
    r0 = p0.add_run(title)
    _set_run_fonts(r0, FONT_HEADING_ZH, FONT_HEADING_EN, PT_COVER_TITLE, bold=True)
    # 副标题：分析人、分析时间
    lines = []
    if analyst:
        lines.append(f"分析人：{analyst}")
    lines.append(f"分析时间：{report_date}")
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = 1
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _set_run_fonts(r, FONT_BODY_ZH, FONT_BODY_EN, PT_COVER_SUB, bold=False)
    # 分页，正文从下一页开始
    p_break = doc.add_paragraph()
    run_break = p_break.add_run()
    run_break.add_break(WD_BREAK.PAGE)


def _set_run_fonts(run, font_zh: str, font_en: str | None = None, size_pt: float | None = None, bold: bool = False, italic: bool = False):
    """设置 run 字体（output_define §2）：中文用 font_zh，西文用 font_en（若提供）。"""
    run.font.name = font_zh
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if font_en and run._element.rPr is not None:
        rPr = run._element.rPr
        try:
            if hasattr(rPr, "rFonts_ascii"):
                rPr.rFonts_ascii = font_en
            if hasattr(rPr, "rFonts_hAnsi"):
                rPr.rFonts_hAnsi = font_en
            if hasattr(rPr, "rFonts_eastAsia"):
                rPr.rFonts_eastAsia = font_zh
        except Exception:
            pass


def _style_heading1(p):
    """一级标题（output_define §2、§3）：微软雅黑 Bold / Arial Bold，16–18pt，段前 24–30pt、段后 12pt，单倍行距。"""
    p.paragraph_format.space_before = Pt(SPACE_BEFORE_H1)
    p.paragraph_format.space_after = Pt(SPACE_AFTER_H1)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        _set_run_fonts(r, FONT_HEADING_ZH, FONT_HEADING_EN, PT_HEADING1, bold=True)


def _style_heading2(p):
    """二级标题（output_define §2、§3）：微软雅黑 Bold / Arial Bold，14pt，段前 12–18pt、段后 6pt。"""
    p.paragraph_format.space_before = Pt(SPACE_BEFORE_H2)
    p.paragraph_format.space_after = Pt(SPACE_AFTER_H2)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        _set_run_fonts(r, FONT_HEADING_ZH, FONT_HEADING_EN, PT_HEADING2, bold=True)


def _style_heading3(p):
    """三级标题（output_define §2、§3）：微软雅黑 Bold，12pt，段前 6–12pt、段后 3pt。"""
    p.paragraph_format.space_before = Pt(SPACE_BEFORE_H3)
    p.paragraph_format.space_after = Pt(SPACE_AFTER_H3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        _set_run_fonts(r, FONT_HEADING_ZH, FONT_HEADING_EN, PT_HEADING3, bold=True)


def _style_body(p):
    """正文（output_define §2、§3）：微软雅黑 / Calibri，五号 10.5pt，首行缩进 0.74cm，行距 1.25 倍。"""
    p.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING_BODY
    for r in p.runs:
        _set_run_fonts(r, FONT_BODY_ZH, FONT_BODY_EN, PT_BODY, bold=False)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        _style_heading1(p)
    elif level == 2:
        _style_heading2(p)
    elif level == 3:
        _style_heading3(p)


def _is_md_table_separator(line: str) -> bool:
    """判断是否为 Markdown 表格分隔行，如 |---|:---:|---|"""
    line = line.strip()
    if not line.startswith("|") or not line.endswith("|"):
        return False
    inner = line[1:-1]
    return bool(re.match(r"^[\s\-:|]+$", inner))


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    """解析 Markdown 表格行，返回 [[cell, ...], ...]。"""
    rows = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line.startswith("|") or not line.endswith("|"):
            break
        if i == 1 and _is_md_table_separator(line):
            continue
        parts = [p.strip() for p in line.split("|")]
        if parts and parts[0] == "":
            parts = parts[1:]
        if parts and parts[-1] == "":
            parts = parts[:-1]
        if parts:
            rows.append(parts)
    return rows


def _is_md_table_block(lines: list[str]) -> bool:
    """判断一组行是否为 Markdown 表格块。"""
    if len(lines) < 2:
        return False
    if not all(line.strip().startswith("|") and line.strip().endswith("|") for line in lines):
        return False
    return _is_md_table_separator(lines[1].strip())


def _add_table(doc, rows: list[list[str]]):
    """将解析后的表格添加到 docx，并应用 output_define 样式。"""
    if not rows:
        return
    num_rows = len(rows)
    num_cols = max(len(r) for r in rows) if rows else 0
    if num_cols == 0:
        return
    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")
    for i, row_cells in enumerate(rows):
        for j, cell_text in enumerate(row_cells):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                for p in cell.paragraphs:
                    for r in p.runs:
                        _set_run_fonts(r, FONT_BODY_ZH, FONT_BODY_EN, PT_BODY, bold=(i == 0))
    table.allow_autofit = False
    for col in table.columns:
        col.width = Cm(4.0)


def _is_numlist_block(lines: list[str]) -> bool:
    """判断是否为 Markdown 编号列表（1. 2. 或 1) 2)）。"""
    if len(lines) < 1:
        return False
    return all(re.match(r"^\d+[\.\)]\s+", ln) for ln in lines)


def _parse_numlist_items(lines: list[str]) -> list[str]:
    """从编号列表行中提取条目内容，去掉 '1. ' 或 '1) ' 前缀。"""
    items = []
    for ln in lines:
        m = re.match(r"^\d+[\.\)]\s+(.*)", ln)
        items.append(m.group(1).strip() if m else ln.strip())
    return items


def _split_ref_item(text: str) -> tuple[str | None, str]:
    """解析引用条目：若含 ' -- '、'—' 或 '：'，拆为 (文件名/URL, 说明)，否则返回 (None, 全文)。"""
    for sep in (" -- ", "—", "："):
        if sep in text:
            parts = text.split(sep, 1)
            if len(parts) == 2 and parts[0].strip():
                return (parts[0].strip(), parts[1].strip())
    return (None, text)


def _add_numlist_item(doc, item: str, as_ref: bool):
    """添加一条编号列表项；as_ref 时用宋体、悬挂缩进，并支持「文件名 -- 说明」加粗文件名。"""
    p = doc.add_paragraph()
    try:
        p.style = "List Number"
    except Exception:
        pass
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = LINE_SPACING_BODY
    p.paragraph_format.left_indent = Cm(FIRST_LINE_INDENT_CM)
    p.paragraph_format.first_line_indent = Cm(-0.74)
    font_zh = FONT_REF_ZH if as_ref else FONT_BODY_ZH
    font_en = FONT_REF_EN if as_ref else FONT_BODY_EN
    prefix, rest = _split_ref_item(item)
    if prefix:
        r0 = p.add_run(prefix + "  ")
        _set_run_fonts(r0, font_zh, font_en, PT_BODY, bold=True)
        if rest:
            r1 = p.add_run(rest)
            _set_run_fonts(r1, font_zh, font_en, PT_BODY, bold=False)
    else:
        for segment, bold, italic in _parse_inline_format(rest or item):
            if not segment:
                continue
            r = p.add_run(segment)
            _set_run_fonts(r, font_zh, font_en, PT_BODY, bold=bold, italic=italic)


def _add_subheading(doc, text: str):
    """添加小标题（如「本地文件：」「**国内材料**:」），支持 **粗体**、*斜体*，output_define §2 引用/附录。"""
    text = (text or "").strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for segment, bold, italic in _parse_inline_format(text):
        if not segment:
            continue
        r = p.add_run(segment)
        _set_run_fonts(r, FONT_REF_ZH, FONT_REF_EN, PT_SUBHEADING, bold=bold, italic=italic)


def _split_lines_by_headings(lines: list[str]) -> list[tuple]:
    """将多行按 Markdown 标题行拆成 (type, content) 列表，避免 ####/###/## 被当作文本。
    type 为 'h1'|'h2'|'h3'|'para'；para 可能带第三元 as_bullet，此处统一不传。"""
    result = []
    current_para = []
    for line in lines:
        s = line.strip()
        if s.startswith("#### ") and not s.startswith("##### "):
            if current_para:
                result.append(("para", " ".join(current_para)))
                current_para = []
            result.append(("h3", s[5:].strip()))
        elif s.startswith("### ") and not s.startswith("#### "):
            if current_para:
                result.append(("para", " ".join(current_para)))
                current_para = []
            result.append(("h2", s[4:].strip()))
        elif s.startswith("## ") and not s.startswith("### "):
            if current_para:
                result.append(("para", " ".join(current_para)))
                current_para = []
            result.append(("h1", s[3:].strip()))
        else:
            current_para.append(s)
    if current_para:
        result.append(("para", " ".join(current_para)))
    return result


def _parse_inline_format(text: str) -> list[tuple[str, bool, bool]]:
    """将段落中的 **粗体** 与 *斜体* 拆成 (片段, bold, italic) 列表，星号不输出。
    优先解析 **，再在每段内解析单 *；*xxx* 转为斜体（与常见 Markdown 一致）。"""
    result = []
    # 先按 ** 拆成粗体/非粗体段
    by_bold = re.split(r"\*\*", text)
    for i, part in enumerate(by_bold):
        is_bold = i % 2 == 1
        # 在每段内再按单 * 拆（单 * 不能是 ** 的一部分，此处 part 内已无 **）
        by_italic = re.split(r"(?<!\*)\*(?!\*)", part)
        for j, seg in enumerate(by_italic):
            is_italic = j % 2 == 1
            if seg:
                result.append((seg, is_bold, is_italic))
    return result


def _add_paragraph_with_inline(doc, text: str, as_bullet: bool = False):
    """添加一段落：去除首行 '- '，将 **x**、*x* 转为粗体/斜体 run，并应用正文样式。"""
    text = (text or "").strip()
    if not text:
        return
    if text in ("---", "***", "----"):
        return
    if text.startswith("- "):
        text = text[2:].strip()
    p = doc.add_paragraph()
    if as_bullet:
        try:
            p.style = "List Bullet"
        except Exception:
            p.paragraph_format.left_indent = Cm(FIRST_LINE_INDENT_CM)
            p.paragraph_format.first_line_indent = Cm(-0.25)
    else:
        p.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING_BODY
    for segment, bold, italic in _parse_inline_format(text):
        if not segment:
            continue
        r = p.add_run(segment)
        _set_run_fonts(r, FONT_BODY_ZH, FONT_BODY_EN, PT_BODY, bold=bold, italic=italic)
    # 不再调用 _style_body(p)，避免覆盖 ** / * 解析出的格式


def add_para(doc, text, as_bullet: bool = False):
    """添加正文段落；支持 **粗体**、*斜体*、去除首行 '- '、忽略 '---'。"""
    text = (text or "").strip()
    if not text or text in ("---", "***", "----"):
        return
    _add_paragraph_with_inline(doc, text, as_bullet=as_bullet)


def parse_md(md_path: Path) -> list:
    """解析 report_draft.md：返回 [(type, content), ...]，type 为 'h1'|'h2'|'h3'|'para'|'list'|'table'。"""
    text = md_path.read_text(encoding="utf-8")
    segments = re.split(r"\n\s*\n", text)
    blocks = []
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        if seg in ("---", "***", "----"):
            continue
        lines = [ln.strip() for ln in seg.split("\n") if ln.strip()]
        if not lines:
            continue
        first = lines[0]
        rest_lines = lines[1:]

        def _append_rest_as_blocks(rest: list[str], heading_text: str):
            if not rest:
                return
            joined = " ".join(rest)
            if joined in ("---", "***", "----"):
                return
            as_ref = "引用" in heading_text or "附录" in heading_text
            if len(rest) >= 2 and rest[0].endswith(":") and len(rest[0]) <= 30 and _is_numlist_block(rest[1:]):
                blocks.append(("subheading", rest[0]))
                blocks.append(("numlist", _parse_numlist_items(rest[1:]), as_ref))
            elif _is_numlist_block(rest):
                blocks.append(("numlist", _parse_numlist_items(rest), as_ref))
            else:
                for sub in _split_lines_by_headings(rest):
                    blocks.append(sub)

        if first.startswith("## ") and not first.startswith("### ") and not first.startswith("#### "):
            h = first[3:].strip()
            blocks.append(("h1", h))
            _append_rest_as_blocks(rest_lines, h)
        elif first.startswith("### ") and not first.startswith("#### "):
            h = first[4:].strip()
            blocks.append(("h2", h))
            _append_rest_as_blocks(rest_lines, h)
        elif first.startswith("#### "):
            h = first[5:].strip()
            blocks.append(("h3", h))
            _append_rest_as_blocks(rest_lines, h)
        elif first.startswith("# ") and not first.startswith("## "):
            continue
        elif all(ln.startswith("- ") for ln in lines) and len(lines) > 1:
            blocks.append(("list", [ln[2:].strip() for ln in lines]))
        elif len(lines) == 1 and first.startswith("- "):
            blocks.append(("para", first[2:].strip(), True))
        elif _is_numlist_block(lines):
            items = _parse_numlist_items(lines)
            blocks.append(("numlist", items, False))
        elif len(lines) >= 2 and first.endswith(":") and len(first) <= 30 and _is_numlist_block(lines[1:]):
            blocks.append(("subheading", first))
            items = _parse_numlist_items(lines[1:])
            as_ref = "本地文件" in first or "网页" in first or "引用" in first or "URL" in first
            blocks.append(("numlist", items, as_ref))
        elif _is_md_table_block(lines):
            rows = _parse_md_table(lines)
            if rows:
                blocks.append(("table", rows))
        else:
            for sub in _split_lines_by_headings(lines):
                blocks.append(sub)
    return blocks


def main():
    import argparse
    # Windows 终端默认编码可能为 GBK，遇到不可编码字符会导致 print() 崩溃。
    # 这里统一将 stdout/stderr 设为 UTF-8 + replace，确保脚本稳定输出日志。
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    try:
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    p = argparse.ArgumentParser(description="从 report_draft.md 生成符合 output_define 样式的报告 docx")
    p.add_argument("--draft", type=Path, default=None, help="report_draft.md 路径（默认 extracted_content/<竞品>/report_draft.md）")
    p.add_argument("--competitor", type=str, default="Inspur", help="竞品名，用于默认草稿路径、输出文件名及封面主标题")
    p.add_argument("--title", type=str, default=None, help="封面主标题（默认 <竞品名>竞品分析，如 浪潮竞品分析）")
    p.add_argument("--analyst", type=str, default="", help="封面副标题：分析人")
    p.add_argument("--out", type=Path, default=None, help="输出 docx 路径（默认 reports/<竞品>_v1.0_<日期>.docx）")
    p.add_argument("--reports-root", type=Path, default=None, help="reports 所在目录（默认项目根）")
    args = p.parse_args()

    root = _project_root()
    reports_root = (args.reports_root or root).resolve()
    reports_dir = reports_root / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    draft = args.draft
    if draft is None:
        draft = root / "extracted_content" / args.competitor / "report_draft.md"
    draft = draft.resolve()
    if not draft.is_file():
        raise SystemExit(f"草稿不存在: {draft}")

    out = args.out
    if out is None:
        out = reports_dir / f"{args.competitor}_v1.0_{date.today().strftime('%Y%m%d')}.docx"
    out = out.resolve()

    cover_title = (args.title or f"{args.competitor}竞品分析").strip()
    report_date = date.today().strftime("%Y-%m-%d")

    blocks = parse_md(draft)
    doc = Document()
    _apply_page_setup(doc)
    _add_cover_page(doc, cover_title, (args.analyst or "").strip(), report_date)
    last_heading = ""
    for block in blocks:
        kind = block[0]
        content = block[1] if len(block) > 1 else ""
        if kind == "h1":
            last_heading = content
            add_heading(doc, content, level=1)
        elif kind == "h2":
            last_heading = content
            add_heading(doc, content, level=2)
        elif kind == "h3":
            last_heading = content
            add_heading(doc, content, level=3)
        elif kind == "subheading":
            _add_subheading(doc, content)
        elif kind == "list":
            for item in content:
                add_para(doc, item, as_bullet=True)
        elif kind == "numlist":
            as_ref = block[2] if len(block) > 2 else ("引用" in last_heading or "附录" in last_heading)
            for item in content:
                _add_numlist_item(doc, item, as_ref=as_ref)
        elif kind == "table":
            _add_table(doc, content)
        elif kind == "para":
            as_bullet = block[2] if len(block) > 2 else False
            add_para(doc, content, as_bullet=as_bullet)
    doc.save(str(out))
    print("已生成报告:", out)


if __name__ == "__main__":
    main()
