#!/usr/bin/env python3
"""
Step 3 迭代2：根据「图片放置清单」JSON，在目标报告 docx 的指定标题后插入图片并添加图注。
支持两种放置清单格式：
1. 扁平格式：[{"heading": "2.1 架构&技术分析", "image": "path/to/image.png", "caption": "系统架构图"}, ...]
2. 嵌套格式：{"placements": [{"chapter_title": "2.1 架构&技术分析", "images": [{"filename": "xxx.png", "caption": "..."}, ...]}, ...]}
"""
import argparse
import json
from pathlib import Path
import sys

try:
    from docx import Document
except ImportError:
    raise SystemExit("请安装 python-docx：pip install python-docx")

try:
    from docx.shared import Cm, Pt
except Exception:  # pragma: no cover
    Cm = None  # type: ignore
    Pt = None  # type: ignore


def find_paragraph_index(doc, heading: str) -> int | None:
    """在文档中查找包含 heading 的段落索引（完全匹配或 strip 后匹配）。"""
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if t == heading or heading in t:
            return i
    return None


def main():
    # Windows 终端默认编码可能为 GBK，遇到不可编码字符会导致输出失败；统一使用 UTF-8 + replace。
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    try:
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    parser = argparse.ArgumentParser(
        description="按放置清单将图片插入报告 docx 的指定标题后"
    )
    parser.add_argument("--docx", type=Path, required=True, help="目标报告 docx 路径")
    parser.add_argument(
        "--placement",
        type=Path,
        required=True,
        help="放置清单 JSON 路径（支持扁平或嵌套格式，详见脚本顶部注释）",
    )
    parser.add_argument(
        "--width-cm",
        type=float,
        default=14,
        help="插入图片宽度（厘米），默认 14",
    )
    args = parser.parse_args()

    docx_path = args.docx.resolve()
    if not docx_path.is_file():
        raise SystemExit(f"docx 文件不存在: {docx_path}")

    placement_path = args.placement.resolve()
    if not placement_path.is_file():
        raise SystemExit(f"放置清单不存在: {placement_path}")

    with open(placement_path, encoding="utf-8") as f:
        data = json.load(f)
    # 支持两种格式：
    # 1. 顶层数组：[{"heading": ..., "image": ..., "caption": ...}, ...]
    # 2. 对象包裹 + 嵌套：{"placements": [{"chapter_title": ..., "images": [{...}]}, ...]}
    if isinstance(data, list):
        placements = data
        nested = False
    elif isinstance(data, dict) and "placements" in data:
        placements = data["placements"]
        nested = True
    else:
        raise SystemExit("放置清单格式不正确：应为数组或含 'placements' 键的对象")
    if not isinstance(placements, list):
        raise SystemExit("放置清单应为 JSON 数组")

    doc = Document(str(docx_path))
    # 图片路径在放置清单中一般为相对于项目根的路径，或相对于 extracted_content/<competitor>/images
    base_dir = Path.cwd()
    # 尝试推断图片目录（与放置清单同级的 images 目录）
    images_dir = placement_path.parent / "images"
    inserted = 0

    # 展开嵌套格式为扁平列表
    flat_items = []
    if nested:
        for chapter in placements:
            heading = (chapter.get("chapter_title") or "").strip()
            for img in chapter.get("images") or []:
                flat_items.append({
                    "heading": heading,
                    "image": (img.get("filename") or "").strip(),
                    "caption": (img.get("caption") or "").strip(),
                })
    else:
        flat_items = placements

    for item in flat_items:
        heading = (item.get("heading") or "").strip()
        image_rel = (item.get("image") or "").strip()
        caption = (item.get("caption") or "").strip()
        if not heading or not image_rel:
            continue
        # 解析图片路径：优先检查 images 子目录，再回退到相对于 base_dir
        if Path(image_rel).is_absolute():
            image_path = Path(image_rel)
        elif (images_dir / image_rel).is_file():
            image_path = (images_dir / image_rel).resolve()
        else:
            image_path = (base_dir / image_rel).resolve()
        if not image_path.is_file():
            print(f"跳过（图片不存在）: {image_path}", file=__import__("sys").stderr)
            continue
        idx = find_paragraph_index(doc, heading)
        if idx is None:
            print(f"跳过（未找到标题）: {heading}", file=__import__("sys").stderr)
            continue
        if idx + 1 >= len(doc.paragraphs):
            print(f"跳过（标题后无段落）: {heading}", file=__import__("sys").stderr)
            continue
        next_para = doc.paragraphs[idx + 1]
        # 顺序必须为：标题 -> 图片 -> 图注（图在上、图注在下，符合 output_define）。先插图片段再插图注段，均 insert_before(next_para)，故先插的图片在前、后插的图注在图片与正文之间。
        img_para = next_para.insert_paragraph_before()
        img_para.alignment = 1  # 居中
        run = img_para.add_run()
        try:
            if Cm is None:
                raise RuntimeError("docx.shared.Cm not available")
            run.add_picture(str(image_path), width=Cm(args.width_cm))
        except Exception as e:
            print(f"插入图片失败 {image_path}: {e}", file=__import__("sys").stderr)
            continue
        if caption:
            cap_para = next_para.insert_paragraph_before(text=caption)
            cap_para.alignment = 1  # 图注默认居中（与图一致）
            # 尽量按 output_define 的题注风格：小五(9pt)、灰色可选；此处仅设置字号，颜色保持默认黑色以兼容打印。
            if Pt is not None:
                for r in cap_para.runs:
                    try:
                        r.font.size = Pt(9)
                        # 中文字体优先用微软雅黑，与正文一致；西文字体不强制，避免不同环境异常
                        r.font.name = "微软雅黑"
                    except Exception:
                        pass
        inserted += 1

    doc.save(str(docx_path))
    print(f"已插入 {inserted} 张图片，已保存: {docx_path}", file=__import__("sys").stderr)


if __name__ == "__main__":
    main()
